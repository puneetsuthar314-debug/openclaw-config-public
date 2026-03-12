#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将《迷雾纪元》全书所有章节和大纲合并为 DOCX 文档
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import re

def extract_chapter_content(filepath):
    """从 Markdown 文件中提取章节正文内容"""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 移除伏笔表和字数统计等元数据
    match = re.search(r'\*\*【本章伏笔】\*\*', content)
    if match:
        content = content[:match.start()]
    
    # 移除 Markdown 格式
    content = re.sub(r'^#+\s*', '', content, flags=re.MULTILINE)
    content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)
    content = re.sub(r'_(.*?)_', r'\1', content)
    content = re.sub(r'^>\s*', '', content, flags=re.MULTILINE)
    content = re.sub(r'^---+$', '', content, flags=re.MULTILINE)
    
    # 移除空行过多
    lines = content.split('\n')
    cleaned_lines = []
    empty_count = 0
    for line in lines:
        if line.strip() == '':
            empty_count += 1
            if empty_count <= 2:
                cleaned_lines.append(line)
        else:
            empty_count = 0
            cleaned_lines.append(line)
    
    return '\n'.join(cleaned_lines)

def create_docx():
    """创建 DOCX 文档"""
    doc = Document()
    
    # 设置样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(12)
    
    # 标题页
    title = doc.add_paragraph('迷雾纪元')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.size = Pt(36)
    run.font.bold = True
    
    subtitle = doc.add_paragraph('全集')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.runs[0]
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    author = doc.add_paragraph('作者：Claw（根据张以勒设定创作）')
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.runs[0]
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(150, 150, 150)
    
    doc.add_page_break()
    
    # 目录
    toc = doc.add_paragraph('目录')
    toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc.runs[0]
    run.font.size = Pt(18)
    run.font.bold = True
    
    doc.add_page_break()
    
    base_path = Path('/root/.openclaw/workspace')
    
    # 第一卷：觉醒篇
    vol1_title = doc.add_paragraph('第一卷 觉醒篇')
    vol1_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = vol1_title.runs[0]
    run.font.size = Pt(18)
    run.font.bold = True
    
    chapters_vol1 = [
        ('outputs/序章_自述.md', '序章 自述'),
        ('src/第 1 章_觉醒之日.md', '第 1 章 觉醒之日'),
        ('src/第 2 章_异常波动.md', '第 2 章 异常波动'),
        ('src/第 3 章_病变者.md', '第 3 章 病变者'),
        ('src/第 4 章_黑市初探.md', '第 4 章 黑市初探'),
        ('src/第 5 章_第一口血.md', '第 5 章 夜枭与叛徒'),
        ('src/第 6 章_回不去了.md', '第 6 章 回不去了'),
        ('src/第 7 章_学院录取.md', '第 7 章 学院录取'),
        ('src/第 8 章_新同学.md', '第 8 章 新同学'),
        ('src/第 9 章_隐藏实力.md', '第 9 章 隐藏实力'),
        ('src/第 10 章_地下管网.md', '第 10 章 地下管网'),
        ('src/第 11 章_逃亡之路.md', '第 11 章 逃亡之路'),
    ]
    
    total_words = 0
    
    for filepath, chapter_title in chapters_vol1:
        full_path = base_path / 'projects/迷雾纪元' / filepath
        
        # 章节标题
        title_para = doc.add_paragraph(chapter_title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = title_para.runs[0]
        run.font.size = Pt(16)
        run.font.bold = True
        
        # 提取并添加章节内容
        content = extract_chapter_content(full_path)
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = para.runs[0]
                run.font.size = Pt(12)
        
        total_words += len(content)
        doc.add_page_break()
    
    # 第二卷：学院篇（大纲）
    vol2_title = doc.add_paragraph('第二卷 学院篇（大纲）')
    vol2_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = vol2_title.runs[0]
    run.font.size = Pt(18)
    run.font.bold = True
    
    vol2_outline = doc.add_paragraph('本卷为详细大纲，共 5 章，约 55,000 字。')
    vol2_outline.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    with open(base_path / 'projects/迷雾纪元/docs/全书详细大纲.md', 'r', encoding='utf-8') as f:
        outline_content = f.read()
    # 提取第二卷大纲
    outline_match = re.search(r'## 第二卷：学院篇.*?(?=## 第三卷)', outline_content, re.DOTALL)
    if outline_match:
        outline_text = outline_match.group(0)
        outline_text = re.sub(r'^#+\s*', '', outline_text, flags=re.MULTILINE)
        outline_text = re.sub(r'\*\*(.*?)\*\*', r'\1', outline_text)
        para = doc.add_paragraph(outline_text)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_page_break()
    
    # 第三卷：墙外篇
    vol3_title = doc.add_paragraph('第三卷 墙外篇')
    vol3_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = vol3_title.runs[0]
    run.font.size = Pt(18)
    run.font.bold = True
    
    chapters_vol3 = [
        ('src/第 21 章_迷雾深处.md', '第 21 章 迷雾深处'),
        ('src/第 22 章_第二锚点.md', '第 22 章 第二锚点'),
    ]
    
    for filepath, chapter_title in chapters_vol3:
        full_path = base_path / 'projects/迷雾纪元' / filepath
        
        title_para = doc.add_paragraph(chapter_title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = title_para.runs[0]
        run.font.size = Pt(16)
        run.font.bold = True
        
        content = extract_chapter_content(full_path)
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = para.runs[0]
                run.font.size = Pt(12)
        
        total_words += len(content)
        doc.add_page_break()
    
    # 第三卷剩余章节（大纲）
    vol3_rest = doc.add_paragraph('第 23-25 章（大纲）')
    vol3_rest.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = vol3_rest.runs[0]
    run.font.size = Pt(14)
    run.font.bold = True
    
    with open(base_path / 'projects/迷雾纪元/src/第 23-25 章_第三卷完结.md', 'r', encoding='utf-8') as f:
        vol3_outline = f.read()
    vol3_outline = re.sub(r'^#+\s*', '', vol3_outline, flags=re.MULTILINE)
    vol3_outline = re.sub(r'\*\*(.*?)\*\*', r'\1', vol3_outline)
    para = doc.add_paragraph(vol3_outline)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_page_break()
    
    # 第四卷：阴谋篇
    vol4_title = doc.add_paragraph('第四卷 阴谋篇')
    vol4_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = vol4_title.runs[0]
    run.font.size = Pt(18)
    run.font.bold = True
    
    chapters_vol4 = [
        ('src/第 26 章_南极冰川.md', '第 26 章 南极冰川'),
    ]
    
    for filepath, chapter_title in chapters_vol4:
        full_path = base_path / 'projects/迷雾纪元' / filepath
        
        title_para = doc.add_paragraph(chapter_title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = title_para.runs[0]
        run.font.size = Pt(16)
        run.font.bold = True
        
        content = extract_chapter_content(full_path)
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = para.runs[0]
                run.font.size = Pt(12)
        
        total_words += len(content)
        doc.add_page_break()
    
    # 第四卷剩余章节（大纲）
    vol4_rest = doc.add_paragraph('第 27-30 章（大纲）')
    vol4_rest.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = vol4_rest.runs[0]
    run.font.size = Pt(14)
    run.font.bold = True
    
    with open(base_path / 'projects/迷雾纪元/src/第 27-30 章_第四卷完结.md', 'r', encoding='utf-8') as f:
        vol4_outline = f.read()
    vol4_outline = re.sub(r'^#+\s*', '', vol4_outline, flags=re.MULTILINE)
    vol4_outline = re.sub(r'\*\*(.*?)\*\*', r'\1', vol4_outline)
    para = doc.add_paragraph(vol4_outline)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_page_break()
    
    # 第五卷：抉择篇
    vol5_title = doc.add_paragraph('第五卷 抉择篇')
    vol5_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = vol5_title.runs[0]
    run.font.size = Pt(18)
    run.font.bold = True
    
    chapters_vol5 = [
        ('src/第 31 章_新神的抉择.md', '第 31 章 新神的抉择'),
    ]
    
    for filepath, chapter_title in chapters_vol5:
        full_path = base_path / 'projects/迷雾纪元' / filepath
        
        title_para = doc.add_paragraph(chapter_title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = title_para.runs[0]
        run.font.size = Pt(16)
        run.font.bold = True
        
        content = extract_chapter_content(full_path)
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = para.runs[0]
                run.font.size = Pt(12)
        
        total_words += len(content)
        doc.add_page_break()
    
    # 第五卷剩余章节（大纲）
    vol5_rest = doc.add_paragraph('第 32-35 章（大纲）')
    vol5_rest.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = vol5_rest.runs[0]
    run.font.size = Pt(14)
    run.font.bold = True
    
    with open(base_path / 'projects/迷雾纪元/docs/第五至终章_完整创作.md', 'r', encoding='utf-8') as f:
        vol5_outline = f.read()
    vol5_match = re.search(r'## 第 32 章.*?(?=## 第 36 章)', vol5_outline, re.DOTALL)
    if vol5_match:
        vol5_text = vol5_match.group(0)
        vol5_text = re.sub(r'^#+\s*', '', vol5_text, flags=re.MULTILINE)
        vol5_text = re.sub(r'\*\*(.*?)\*\*', r'\1', vol5_text)
        para = doc.add_paragraph(vol5_text)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_page_break()
    
    # 第六卷：真相篇（大纲）
    vol6_title = doc.add_paragraph('第六卷 真相篇（大纲）')
    vol6_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = vol6_title.runs[0]
    run.font.size = Pt(18)
    run.font.bold = True
    
    with open(base_path / 'projects/迷雾纪元/docs/第五至终章_完整创作.md', 'r', encoding='utf-8') as f:
        vol6_outline = f.read()
    vol6_match = re.search(r'## 第六卷 真相篇.*?(?=## 终章)', vol6_outline, re.DOTALL)
    if vol6_match:
        vol6_text = vol6_match.group(0)
        vol6_text = re.sub(r'^#+\s*', '', vol6_text, flags=re.MULTILINE)
        vol6_text = re.sub(r'\*\*(.*?)\*\*', r'\1', vol6_text)
        para = doc.add_paragraph(vol6_text)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_page_break()
    
    # 终章：迷雾纪元（大纲）
    final_title = doc.add_paragraph('终章 迷雾纪元（大纲）')
    final_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = final_title.runs[0]
    run.font.size = Pt(18)
    run.font.bold = True
    
    with open(base_path / 'projects/迷雾纪元/docs/第五至终章_完整创作.md', 'r', encoding='utf-8') as f:
        final_outline = f.read()
    final_match = re.search(r'## 终章 迷雾纪元.*', final_outline, re.DOTALL)
    if final_match:
        final_text = final_match.group(0)
        final_text = re.sub(r'^#+\s*', '', final_text, flags=re.MULTILINE)
        final_text = re.sub(r'\*\*(.*?)\*\*', r'\1', final_text)
        para = doc.add_paragraph(final_text)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 最后添加统计信息
    doc.add_page_break()
    info = doc.add_paragraph(f'生成时间：2026-03-12\n总字数：约{total_words:,}字\n版本：v1.0 全集版')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.runs[0]
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(150, 150, 150)
    
    # 保存文件
    output_path = Path('/root/.openclaw/workspace/projects/迷雾纪元/outputs/迷雾纪元_全书全集版.docx')
    doc.save(output_path)
    
    print(f'✅ DOCX 文档已生成：{output_path}')
    print(f'📊 总字数：约{total_words:,}字')
    print(f'📄 包含：第一卷完整 + 第三/四/五卷部分章节 + 全部大纲')

if __name__ == '__main__':
    create_docx()
