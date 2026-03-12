#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将《迷雾纪元》已创作章节合并为 DOCX 文档
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import re

def extract_chapter_content(filepath):
    """从 Markdown 文件中提取章节正文内容"""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 移除伏笔表和字数统计等元数据
    # 保留从标题到【本章伏笔】之前的内容
    match = re.search(r'\*\*【本章伏笔】\*\*', content)
    if match:
        content = content[:match.start()]
    
    # 移除 Markdown 格式
    content = re.sub(r'^#+\s*', '', content, flags=re.MULTILINE)  # 移除标题标记
    content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)  # 移除粗体标记
    content = re.sub(r'_(.*?)_', r'\1', content)  # 移除斜体标记
    content = re.sub(r'^>\s*', '', content, flags=re.MULTILINE)  # 移除引用标记
    content = re.sub(r'^---+$', '', content, flags=re.MULTILINE)  # 移除分隔线
    
    # 移除空行过多的部分
    lines = content.split('\n')
    cleaned_lines = []
    empty_count = 0
    for line in lines:
        if line.strip() == '':
            empty_count += 1
            if empty_count <= 2:  # 最多保留 2 个连续空行
                cleaned_lines.append(line)
        else:
            empty_count = 0
            cleaned_lines.append(line)
    
    return '\n'.join(cleaned_lines)

def create_docx():
    """创建 DOCX 文档"""
    doc = Document()
    
    # 设置样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(12)
    
    # 标题
    title = doc.add_paragraph('迷雾纪元')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # 副标题
    subtitle = doc.add_paragraph('第一卷：觉醒篇')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.runs[0]
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()  # 空行
    
    # 章节文件列表
    base_path = Path('/root/.openclaw/workspace')
    chapters = [
        (base_path / 'projects/迷雾纪元/outputs/序章_自述.md', '序章 自述'),
        (base_path / 'projects/迷雾纪元/src/第 1 章_觉醒之日.md', '第 1 章 觉醒之日'),
        (base_path / 'projects/迷雾纪元/src/第 2 章_异常波动.md', '第 2 章 异常波动'),
        (base_path / 'projects/迷雾纪元/src/第 3 章_病变者.md', '第 3 章 病变者'),
        (base_path / 'projects/迷雾纪元/src/第 4 章_黑市初探.md', '第 4 章 黑市初探'),
        (base_path / 'projects/迷雾纪元/src/第 5 章_第一口血.md', '第 5 章 夜枭与叛徒'),
        (base_path / 'projects/迷雾纪元/src/第 6 章_回不去了.md', '第 6 章 回不去了'),
        (base_path / 'projects/迷雾纪元/src/第 7 章_学院录取.md', '第 7 章 学院录取'),
        (base_path / 'projects/迷雾纪元/src/第 8 章_新同学.md', '第 8 章 新同学'),
        (base_path / 'projects/迷雾纪元/src/第 9 章_隐藏实力.md', '第 9 章 隐藏实力'),
        (base_path / 'projects/迷雾纪元/src/第 10 章_地下管网.md', '第 10 章 地下管网'),
        (base_path / 'projects/迷雾纪元/src/第 11 章_逃亡之路.md', '第 11 章 逃亡之路'),
        (base_path / 'projects/迷雾纪元/src/第 21 章_迷雾深处.md', '第 21 章 迷雾深处'),
        (base_path / 'projects/迷雾纪元/src/第 22 章_第二锚点.md', '第 22 章 第二锚点'),
        (base_path / 'projects/迷雾纪元/src/第 26 章_南极冰川.md', '第 26 章 南极冰川'),
        (base_path / 'projects/迷雾纪元/src/第 31 章_新神的抉择.md', '第 31 章 新神的抉择'),
    ]
    
    total_words = 0
    
    for filepath, chapter_title in chapters:
        # 章节标题
        title_para = doc.add_paragraph(chapter_title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = title_para.runs[0]
        run.font.size = Pt(16)
        run.font.bold = True
        
        # 提取并添加章节内容
        content = extract_chapter_content(filepath)
        
        # 分段添加内容
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = para.runs[0]
                run.font.size = Pt(12)
        
        # 添加分页符（除了最后一章）
        if filepath != chapters[-1][0]:
            doc.add_page_break()
        
        # 统计字数（估算）
        total_words += len(content)
    
    # 最后添加统计信息
    doc.add_page_break()
    info_text = f'生成时间：2026-03-12\n总字数：约{total_words:,}字\n版本：v1.0'
    info = doc.add_paragraph(info_text)
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.runs[0]
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(150, 150, 150)
    
    # 保存文件
    output_path = Path('/root/.openclaw/workspace/projects/迷雾纪元/outputs/迷雾纪元_全书已创作章节.docx')
    doc.save(output_path)
    
    print(f'✅ DOCX 文档已生成：{output_path}')
    print(f'📊 总字数：约{total_words:,}字')
    print(f'📄 包含章节：序章 + 第 1-11 章 + 第 21-22 章 + 第 26 章 + 第 31 章（第一/三/四/五卷开篇）')

if __name__ == '__main__':
    create_docx()
